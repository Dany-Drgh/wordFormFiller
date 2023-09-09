from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd


document = Document("ARR.docx")



vip_country_table = document.tables[0]
row = vip_country_table.rows[1]
vip_name_title_cell = row.cells[0]
country_cell = row.cells[2]

flight_info_table = document.tables[1]
row = flight_info_table.rows[1]
date_cell = row.cells[2]
time_cell = row.cells[4]
flight_cell = row.cells[6]

input_vip_name = input("Please, enter VIP name and title \n> ")

addPax = input("Are there any additional passenger ? (Y/n)\n> ") or "Y"

if addPax == "Y" or addPax == "y":
    input_add_pax = input("Please enter the number of additional passenger:\n> ")
    if input_add_pax != "0":
        vip_name_title_cell.text = input_vip_name +" + " + input_add_pax + " Pax."
    else:
        vip_name_title_cell.text = input_vip_name
else:
    vip_name_title_cell.text = input_vip_name

input_country = input("Please enter the country on behalf of which this request is being filled:\n> ")
country_cell.text = input_country
country_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

input_date = input("Please enter the date of the flight:\n> ")
date_cell.text = input_date
date_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

input_time = input("Please enter the time of the flight:\n> ")
time_cell.text = input_time
time_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

input_flight = input("Please enter the flight number:\n> ")
flight_cell.text = input_flight
flight_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER



input_limo_file_path = input("Enter the path of the excel file containing the list of limousines:\n> ")
limo_df = pd.read_excel(input_limo_file_path)

limo_table = document.tables[3]

if len(limo_df) > 15:
    print("There are more than 15 limousines in the list. Which is not allowed GVA \nPlease check the list and try again.")
    exit()

for i in range(0, len(limo_df)):

    row = limo_table.rows[i+1]
    plate_cell = row.cells[2]
    driver_last_name_cell = row.cells[3]
    driver_first_name_cell = row.cells[4]
    driver_DOB_cell = row.cells[5]
    driver_Nationality_cell = row.cells[6]

    plate_cell.text = limo_df.iloc[i,0]
    driver_last_name_cell.text = limo_df.iloc[i,1]
    driver_first_name_cell.text = limo_df.iloc[i,2]
    driver_DOB_cell.text = limo_df.iloc[i,3].strftime("%d/%m/%Y")
    driver_Nationality_cell.text = limo_df.iloc[i,4]

print("Limo list is processed!")


input_truck_file_path = input("Enter the path of the excel file containing the list of Luggage Trucks:\n> ")
truck_df = pd.read_excel(input_truck_file_path)

truck_table = document.tables[3]

if len(truck_df) > 6:
    print("There are more than 6 limousines in the list. Which is not allowed in this script for now but we are working on it.. \nPlease check the list and try again.")
    exit()

for i in range(0, len(truck_df)):

    row = limo_table.rows[i+17]
    plate_cell = row.cells[2]
    driver_last_name_cell = row.cells[3]
    driver_first_name_cell = row.cells[4]
    driver_DOB_cell = row.cells[5]
    driver_Nationality_cell = row.cells[6]

    plate_cell.text = limo_df.iloc[i,0]
    driver_last_name_cell.text = limo_df.iloc[i,1]
    driver_first_name_cell.text = limo_df.iloc[i,2]
    driver_DOB_cell.text = limo_df.iloc[i,3].strftime("%d/%m/%Y")
    driver_Nationality_cell.text = limo_df.iloc[i,4]

print("Truck list is processed!")



print("Saving file...")
document.save("ARR "+input_vip_name+" "+input_flight+" "+ input_date.replace("/","-")+".docx")
document.save("ARR "+input_vip_name+" "+input_flight+" "+ input_date.replace("/","-")+".pdf")

print("File saved! Please check your file. ")